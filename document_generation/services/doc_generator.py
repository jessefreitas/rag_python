"""
Gerador de Documentos Jurídicos
Integrado com CrewAI e sistema RAG Python
"""

import logging
import os
import uuid
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
from jinja2 import Environment, FileSystemLoader, select_autoescape

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..models.document_models import DocRequest, DocResponse, DocumentTemplate
from privacy_system import PrivacyCompliance
from llm_providers import llm_manager

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Gerador principal de documentos jurídicos"""
    
    def __init__(self):
        self.templates_dir = Path("document_generation/templates")
        self.output_dir = Path("static/documentos_gerados")
        self.privacy_manager = PrivacyCompliance()
        
        # Criar diretórios se não existirem
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Configurar Jinja2
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=select_autoescape(['html', 'xml'])
        )
        
        # Templates disponíveis
        self.available_templates = {
            'contrato_prestacao_servicos': 'Contrato de Prestação de Serviços',
            'peticao_inicial_civel': 'Petição Inicial Cível',
            'procuracao': 'Procuração',
            'nda_confidencialidade': 'Termo de Confidencialidade',
            'parecer_juridico': 'Parecer Jurídico',
            'contestacao': 'Contestação',
            'recurso_apelacao': 'Recurso de Apelação'
        }
    
    def generate_document(self, request: DocRequest) -> DocResponse:
        """Gera um documento baseado na requisição"""
        
        try:
            logger.info(f"Iniciando geração de documento: {request.tipo_documento}")
            
            # Verificar privacidade se solicitado
            if hasattr(request, 'check_privacy') and request.check_privacy:
                privacy_check = self._check_document_privacy(request)
                if not privacy_check['approved']:
                    return DocResponse(
                        status="erro",
                        mensagem=f"Falha na verificação de privacidade: {', '.join(privacy_check['issues'])}",
                        agent_id=request.agent_id
                    )
            
            # Melhorar conteúdo com IA se solicitado
            if request.use_ai_enhancement:
                request = self._enhance_with_ai(request)
            
            # Gerar documento
            doc_path = self._generate_docx(request)
            
            # Converter para PDF se solicitado
            if request.formato == "pdf":
                from .pdf_converter import PDFConverter
                converter = PDFConverter()
                doc_path = converter.convert_to_pdf(doc_path)
            
            # Criar resposta
            file_size = os.path.getsize(doc_path)
            file_name = os.path.basename(doc_path)
            
            response = DocResponse(
                status="sucesso",
                mensagem="Documento gerado com sucesso",
                nome_arquivo=file_name,
                url_arquivo=f"/static/documentos_gerados/{file_name}",
                tipo_documento=request.tipo_documento,
                formato=request.formato,
                tamanho_bytes=file_size,
                agent_id=request.agent_id,
                ai_enhancement_used=request.use_ai_enhancement
            )
            
            if request.use_ai_enhancement:
                response.ai_model_used = llm_manager.active_provider
            
            logger.info(f"Documento gerado com sucesso: {file_name}")
            return response
            
        except Exception as e:
            logger.error(f"Erro na geração do documento: {e}")
            return DocResponse(
                status="erro",
                mensagem=f"Erro na geração: {str(e)}",
                agent_id=request.agent_id
            )
    
    def _generate_docx(self, request: DocRequest) -> str:
        """Gera documento DOCX"""
        
        # Verificar se template existe
        template_name = f"{request.tipo_documento}.jinja2"
        template_path = self.templates_dir / template_name
        
        if not template_path.exists():
            # Criar template básico se não existir
            self._create_basic_template(request.tipo_documento)
        
        # Carregar template
        template = self.jinja_env.get_template(template_name)
        
        # Renderizar conteúdo
        rendered_content = template.render(**request.variaveis)
        
        # Criar documento Word
        doc = Document()
        
        # Configurar estilos
        self._setup_document_styles(doc)
        
        # Adicionar conteúdo
        self._add_content_to_document(doc, rendered_content, request)
        
        # Salvar documento
        file_name = f"{request.tipo_documento}_{uuid.uuid4().hex[:8]}.docx"
        file_path = self.output_dir / file_name
        
        doc.save(str(file_path))
        
        return str(file_path)
    
    def _setup_document_styles(self, doc: Document):
        """Configura estilos do documento"""
        
        # Estilo para título
        try:
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Inches(0.2)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # Estilo já existe
        
        # Estilo para corpo do texto
        try:
            body_style = doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Arial'
            body_style.font.size = Inches(0.15)
            body_style.paragraph_format.line_spacing = 1.5
        except:
            pass  # Estilo já existe
    
    def _add_content_to_document(self, doc: Document, content: str, request: DocRequest):
        """Adiciona conteúdo ao documento"""
        
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detectar se é título (linhas em maiúsculo ou com padrões específicos)
            if line.isupper() or line.startswith('CONTRATO') or line.startswith('PETIÇÃO'):
                p = doc.add_paragraph(line)
                try:
                    p.style = 'CustomTitle'
                except:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
            else:
                p = doc.add_paragraph(line)
                try:
                    p.style = 'CustomBody'
                except:
                    pass
        
        # Adicionar assinatura
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        doc.add_paragraph(f"Local e data: ________________")
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        doc.add_paragraph("Assinatura")
    
    def _enhance_with_ai(self, request: DocRequest) -> DocRequest:
        """Melhora o conteúdo usando IA"""
        
        try:
            # Prompt para melhoramento
            enhancement_prompt = f"""
            Você é um especialista em redação jurídica. Melhore e complete as seguintes variáveis 
            para um documento do tipo "{request.tipo_documento}":
            
            Variáveis atuais:
            {request.variaveis}
            
            Tarefas:
            1. Complete informações faltantes de forma profissional
            2. Melhore a linguagem jurídica
            3. Adicione cláusulas padrão apropriadas
            4. Mantenha conformidade legal
            
            Retorne apenas as variáveis melhoradas em formato JSON.
            """
            
            # Obter resposta da IA
            ai_response = llm_manager.get_response(
                enhancement_prompt,
                temperature=0.3,
                max_tokens=1500
            )
            
            # Tentar parsear resposta como JSON
            import json
            try:
                enhanced_vars = json.loads(ai_response)
                # Mesclar variáveis originais com melhoradas
                request.variaveis.update(enhanced_vars)
                logger.info("Conteúdo melhorado com IA")
            except json.JSONDecodeError:
                logger.warning("Não foi possível parsear resposta da IA, usando variáveis originais")
            
        except Exception as e:
            logger.error(f"Erro no melhoramento com IA: {e}")
        
        return request
    
    def _check_document_privacy(self, request: DocRequest) -> Dict[str, Any]:
        """Verifica aspectos de privacidade do documento"""
        
        # Verificar variáveis
        all_content = " ".join(str(v) for v in request.variaveis.values())
        
        return self.privacy_manager.detect_personal_data_only(all_content, detailed=True)
    
    def _create_basic_template(self, document_type: str):
        """Cria um template básico se não existir"""
        
        templates = {
            'contrato_prestacao_servicos': '''
CONTRATO DE PRESTAÇÃO DE SERVIÇOS

Pelo presente instrumento particular, de um lado {{ contratante_nome }}, 
doravante denominado(a) CONTRATANTE, e de outro lado {{ contratado_nome }}, 
doravante denominado(a) CONTRATADO(a), têm entre si justo e contratado o seguinte:

CLÁUSULA 1ª – OBJETO
O objeto do presente contrato é {{ objeto }}.

CLÁUSULA 2ª – REMUNERAÇÃO
Pela execução dos serviços, o CONTRATANTE pagará ao CONTRATADO(a) o valor de {{ valor }}.

CLÁUSULA 3ª – PRAZO
O presente contrato terá vigência de {{ prazo }}.

CLÁUSULA 4ª – RESCISÃO
O presente contrato poderá ser rescindido por qualquer das partes mediante aviso prévio de {{ aviso_previo }}.

{{ clausulas_extras }}

E, por estarem assim justos e contratados, firmam o presente.
            ''',
            'peticao_inicial_civel': '''
EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {{ vara }} DA COMARCA DE {{ comarca }}

{{ autor_nome }}, {{ autor_qualificacao }}, vem, respeitosamente, à presença de Vossa Excelência, 
por meio de seu advogado que esta subscreve, propor

AÇÃO {{ tipo_acao }}

em face de {{ reu_nome }}, {{ reu_qualificacao }}, pelos fatos e fundamentos jurídicos a seguir expostos:

I - DOS FATOS

{{ fatos }}

II - DO DIREITO

{{ fundamentos_juridicos }}

III - DOS PEDIDOS

Diante do exposto, requer-se:

{{ pedidos }}

Dá-se à causa o valor de {{ valor_causa }}.

Termos em que,
Pede deferimento.

{{ local }}, {{ data }}.

{{ advogado_nome }}
OAB/{{ oab_estado }} {{ oab_numero }}
            ''',
            'procuracao': '''
PROCURAÇÃO

Pelo presente instrumento particular, {{ outorgante_nome }}, {{ outorgante_qualificacao }}, 
nomeia e constitui seu bastante procurador {{ outorgado_nome }}, {{ outorgado_qualificacao }}, 
para o fim especial de {{ finalidade }}, podendo para tanto:

{{ poderes }}

A presente procuração é válida até {{ validade }}.

{{ local }}, {{ data }}.

________________________________
{{ outorgante_nome }}
Outorgante
            '''
        }
        
        if document_type in templates:
            template_path = self.templates_dir / f"{document_type}.jinja2"
            with open(template_path, 'w', encoding='utf-8') as f:
                f.write(templates[document_type].strip())
            
            logger.info(f"Template básico criado para {document_type}")
    
    def list_available_templates(self) -> List[Dict[str, str]]:
        """Lista templates disponíveis"""
        
        templates = []
        for template_id, description in self.available_templates.items():
            template_path = self.templates_dir / f"{template_id}.jinja2"
            templates.append({
                'id': template_id,
                'description': description,
                'exists': template_path.exists(),
                'path': str(template_path)
            })
        
        return templates
    
    def get_template_variables(self, template_id: str) -> List[str]:
        """Obtém variáveis necessárias para um template"""
        
        template_path = self.templates_dir / f"{template_id}.jinja2"
        
        if not template_path.exists():
            return []
        
        # Ler template e extrair variáveis
        with open(template_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extrair variáveis usando regex
        import re
        variables = re.findall(r'\{\{\s*(\w+)\s*\}\}', content)
        
        return list(set(variables))  # Remover duplicatas

# Instância global
document_generator = DocumentGenerator() 